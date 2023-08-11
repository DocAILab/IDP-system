"""
基础函数库
"""

import os
import json
import pathlib

import docx
import numpy as np
import cv2
import paddleocr
from paddleocr import PPStructure,save_structure_res
import pdfplumber
import fitz
import re
from paddlenlp import Taskflow
from pprint import pprint
import jieba
import paddle
import paddle.nn.functional as F
from openpyxl import Workbook
import time
import pickle
import random
import string
from numpy import array, uint64
import logging
import math
import multiprocessing
from utils import *
import pathlib

import difflib
from multiprocessing import Pool, Manager
import hashlib
from nltk import ngrams
from datasketch import MinHash, MinHashLSH
import time


class FileHash:
    def __init__(self):
        # domain_pred_res: 领域预测结果，str
        # domain_pred_match_num: 领域预测结果的置信度，float
        # words: 文件名 / 标题关键词提取结果，set 其中不包括跨领域词
        # domain_matched: 各领域下匹配到的文本关键词，dict {领域: [文本关键词, 匹配的最佳领域词, 相似度]} 其中包括了非预测领域的结果
        # use_other 是否使用了other
        # file_key_word 存储schema和对应schema的结果 为string:list list中为提取出的信息
        self.domain_pred = ""
        self.domain_pred_match_num = 0.0
        self.words = set()
        self.domain_matched = {}
        self.use_other = False
        self.file_key_word = dict()

        self.hash1 = generate_ngram_lsh_fingerprint("hash1")
        self.hash2 = generate_ngram_lsh_fingerprint("hash2")
        # dict {'领域':[[finger1,finger2],[finger1,finger2],[finger1,finger2]]}
        self.hash3 = {}
        self.hash3All = generate_ngram_lsh_fingerprint("hash3")


# schema字典，格式：{领域: {schema类型: [具体schema关键词]}}
# 其中，'others'对应领域分类未知时的情况，保持为空，或可添加通用敏感词
schemas_dict = {
    'industry': {
        'gongchengbaogao': ['工程/项目名称', '工程/项目编号', '实施机构','建设单位','设计单位', '勘察单位','工程/项目地点', '图件名称', '工程负责人/总工程师', '法人','审定人', '审核人', '勘察人', '校对人','审查机构', '勘察人', '工程类别', '项目分类','工程用途', '联系人', '联系人电话'],
        'jingjia': ['项目名称', '负责人', '审核人', '法人', '联系方式', '联系人', '报价有效期限', '姓名', '单价', '总价', '价格', '注册资本', '法定代表人', '地址', '基础软件服务'],
        'jianyan': ['委托方', '生产方', '样品名称', '商标', '检验类型', '判定依据', '使用单位', '设计单位', '建设单位', '测试单位'],
        'shebei': ['使用单位', '样品类别', '型号规格', '检验类型', '设备号', '登记证'],
        'yuanshu': ['托运人', '收获人', '装货人', '承运人', '目的地', '起运地', '许可证号', '身份证', '起运日期', '联系电话', '驾驶员', '押运员', '姓名', '车牌号码', '道路运行证', '从业资格证', '单位名称', '调度人'],
        'biaoshu': ['投标人', '工程项目名称', '组织机构代码', '社会统一信用代码', '项目负责人', '投标保证金', '工期', '解密时间', '中标企业', '中标人', '中标价', '管理费','预算/工程造价'],
        'xiangmu': ['责任单位', '项目名称', '投标人', '投标人地址', '注册地址', '企业名称', '企业注册类型', '传真', '投标报价', '总投资', '交易编号', '营业收入', '利润','管理费', '预算'],
        'fapiao': ['单位名称', '地址', '发票代码', '发票号码', '纳税人识别号', '价税合计', '收款人', '开票人', '付款账号', '收款账号', '付款户名', '收款户名', '付款开户行', '收款开户行', '交易金额'],
        'contract': ['甲方', '乙方', '产品名称', '项目名称', '工程名称', '工程地点', '合同编号', '发包人', '设计人', '委托人', '咨询人', '勘察人', '项目费用/投资金额', '组织机构代码', '纳税人识别号', '开户银行', '银行账号', '地址', '公司', '法人', '代理人', '联络人', '建设单位']},
    'personal': {
        'personal': ['姓名','性别', '年龄', '身份证号', '电话', '邮箱', '总价', '开户行', '银行账号', '单位']},
    'enterprise': {
        'xiangmu': ['责任单位', '项目名称', '投标人', '投标人地址', '注册地址', '传真', '投标报价', '交易编号', '营业收入', '利润', '管理费', '预算'],
        'fapiao': ['单位名称', '地址', '发票代码', '发票号码', '纳税人识别号', '价税合计', '收款人', '开票人'],
        'contract': ['甲方', '乙方', '产品名称', '项目名称', '总价', '金额', '纳税人识别号', '开户银行', '银行账号', '地址', '开户行', '公司', '法人', '代理人', '联络人', '购货方', '供货单位']},
    'hospital': {
        'person': ['健康卡号', '电子邮箱', '家庭住址', '工作单位', '缴费记录', '医保卡号', '居住证', '既往病史', '诊断', '科别', '血型', '过敏源', '残疾情况'],
        'mg': ['收入', '保险', '协议机构', '药品', '用药信息', '病程记录', '手术记录', '病历', '病例', '医嘱', '医疗器械', '临床研究', '医院运营', '传染病', '疾病监测', '医疗交易', '卫生事件名称', '卫生事件时间', '卫生事件地点'],
        'contract': ['甲方', '乙方', '产品名称', '项目名称', '总价', '金额', '纳税人识别号', '开户银行', '银行账号', '地址', '开户行', '公司', '法人', '代理人', '联络人', '购货方', '供货单位']},
    'sechool': {
        'students_information': ['民族','港澳台侨','国家地区','手机号','电子信箱','通讯地址','家庭电话','家庭住址','家庭邮政编码','家庭成员数量','联系方式','父亲姓名','母亲姓名'],
        'students_graduation_information': ['学籍','学籍异动','毕业中学','生源地','学历码','所学专业','入学年月','学制','毕业证号','学位证号','离校状态','档案接受地址'],
        'teacher_information': ['出生日期','最高学历','最高学位','最高学历院校','资格种类','证书号码','认证机构','职务名称','职务级别','资格时间','评审类型','评审学科','学院单位','专业','年工资收入','基本工资','绩效工资','生活补助','其他津贴补贴','五险一金'],
        'finance': ['金额','税务类型','抬头','票本序号','票类码','规格','用途','领用人','联系方法','领用部门号','入帐科目及项目','金额限制','领用日期','归还日期','实开金额'],
        'study_contract': ['项目名称','项目负责人','项目编号','立项日期','开始日期','计划结项日期','结项日期','项目性质','项目分类','位次','完成情况','项目来源','合同编号','合同名称','甲方','乙方','两方','合同类型码','合同金额','终止日期','免证营业税']},
    'government': {
        'project_review': ['项目名称', '项目代码', '建设单位', '项目将设依据', '拟选位置', '规划选址', '经办人', '签发人', '日期','已完成投资', '投资额完成率', '实到资金利用率', '县级配套到位率', '年度计划开工项目数', '已完工项目数', '完工率'],
        'file_review': ['来文单位', '来文字号' ,'密级', '收文日期', '标题', '拟办意见', '领导', '阅件人'],
        'contract': ['甲方', '乙方', '产品名称', '项目名称', '总价', '金额', '纳税人识别号', '开户银行', '银行账号', '地址', '开户行', '公司', '法人', '代理人', '联络人', '购货方', '供货单位']},
    'transport': {
        'entity': ['托运人', '收获人', '装货人', '承运人', '目的地', '起运地', '许可证号', '身份证', '起运日期', '联系电话', '驾驶员', '押运员', '姓名', '车牌号码', '道路运行证', '从业资格证', '单位名称', '调度人'],
        'goods': ['货物编号', '危险货物编号','危险货物类别', '条目类别', '包装类别', '应急措施', '联合国编号', '分类代码', '标志', '特殊规定', '有限数量', '特殊规定', '罐体代码', '罐体编号', '罐体运输车俩', '危险性识别号'],
        'baodan': ['被保险人', '赔偿限额', '死亡伤残赔偿限额', '医疗费用赔偿限额', '财产损失赔偿限额', '保险费', '纳税人识别号', '合计', '投保确认码', '公司名称', '公司地址', '保险单号'],
        'contract': ['甲方', '乙方', '产品名称', '项目名称', '总价', '金额', '纳税人识别号', '开户银行', '银行账号', '地址', '开户行', '公司', '法人', '代理人', '联络人', '购货方', '供货单位']},
    'others': {},
}


sensitive_words = {'秘密', '此件不公开','不得擅自扩散','机密','保密','版权所有','复制','传播','外传'}

schemas_dict_education = {
    '人员域': {
        # 学生身份信息
        '学生-基本信息': ['姓名','姓名拼音','曾用名','民族','性别名称','港澳台侨','国家地区','手机号','电子信箱','微信号','QQ号','通讯地址','邮政编码'],
        '学生-身份信息': ['身份证件类型','身份证件号'],
        '体质健康信息': ['身高等级','体重等级','肺活量等级','仰卧起坐等级','跑步等级','视力等级','篮球等级','足球等级'],
        '家庭成员信息': ['工作单位','家庭电话','家庭住址','家庭邮政编码','家庭关系','家庭成员姓名','联系方式','家长出生日期'],
        '资助信息': ['资助类型','资助年份','资助金额','资助学年','银行卡号','开户行'],
        '家庭经济信息': ['劳动力人口数','年人均收入','是否困难户'],
        # zy
        '学生-网络密码及关联信息': ['系统登录账号','密码'],
        # 学生校内信息
        '学生入学信息':['入学专业','学科门类','学习方式','培养方式','学生类别','是否在职生','在职工号','入学年月','入学年级','现在专业','研究方向','学制','学位类型','学籍状态','是否休学','是否请假','预计毕业年份','预计毕业月份','实际毕业时间'],
        '学籍信息': ['学籍相关信息','学籍异动信息','毕业中学','生源地'],
        '学历学位信息': ['学历码','所学专业','入学年月','学习形式码学制','结束学业年月','毕肄业学校或单位','毕业证号','毕业证书印制号','获得学位码'],
        '档案信息': ['异动日期','异动文号','异动类别','异动原因描述','复学日期','原系所','原班级','原专业','离校状态','档案接受地址','档案去向地','档案邮寄地址','邮寄方式','运单号','原导师','档案邮寄时间'],
        '处分信息': ['违纪类型','违纪日期','处分类型','处分日期','处分原因','处分文号','处分解除日期','处分撤销日期','处分撤销原因','申诉日期','申委会审议日期','申委会审议结论','违纪类别','审委会审议结论'],
        # 教师身份信息
        '教师-基本信息': ['姓名','性别','证件号码','出生日期','民族','最高学历','最高学位层次','获得最高学历的院校或机构'],
        '资格信息': ['资格种类','任教学科','证书号码','认证机构名称','专业技术职务名称','专业技术职务级别','取得资格时间','评审类型','评审学科','是否现职称'],
        '教师-身份信息': ['学工号','学院','单位','专业','班级','导师工号','导师类别','导师类别代码','导师联系方式'],
        '工资待遇信息': ['年工资收入','基本工资','绩效工资','乡村教师生活补助','其他津贴补贴','五险一金'],
        # zy
        '教师-网络密码及关联信息': ['系统登录账号','密码'],
        # 干部身份信息
        '干部-基本信息': ['姓名','性别','证件号码','出生日期','民族','最高学历','最高学位层次','获得最高学历的院校或机构'],
    },
    '财务域': {
        # 学生财务相关
        '一卡通消费信息': ['操作时间','科目描述','钱包交易金额','钱包余额','操作员','工作站','终端名称'],
        '创新基金信息': ['基金级别','资助年限','资助金额','研究成果'],
        # 账务管理
        # zy
        '账务数据': ['类型编号','类型名称','凭证标记','限制科目','限制性质码'],
        '凭证数据': ['凭证编号','分录标记码','凭证标记','凭证摘要','科目编号','部门代码','项目编号','记帐方向码','金额','生成日期','冲暂','支票号','备注','制单','复核','出纳'],
        '经费来源': ['会计年度','入帐月份','凭证标志','凭证日期'],
        '往来账数据': ['金额','原因'],
        '暂存款': ['金额','原因'],
        '暂付款': ['金额','原因'],
        '税务信息': ['金额','税务类型','抬头'],
        '票据信息': ['票本序号','启始号','截止号','票类码','规格','用途','领用人','联系方法','领用部门号','入帐科目及项目','金额限制','领用日期','归还日期','已用截止号','实开金额','有效张数','作废张数','交回人号','验收人号','备注'],
    },
    '资产域': {
        # 实验室设施数据
        '实验室资源数据': ['实验室号','实验室名称','建立日期','建筑面积','使用面积','实验室位置','负责人教工号','管理级别码','实验室类别码','实验室电话','器材管理员教工号','实验室介绍'],
        '实验室仪器': ['仪器号','仪器名称','英文名称','学校单位层次码','产权及使用状况码','分类号码型号','出厂日期','生产国别码','经费科目码','购置日期','厂家','出厂号','单据号'],
        # zy
        '实验室运行': ['实验室号','年度','投入额支出额','年耗电量','年耗水量','维护费'],
    },
    '校园生活': {
        # 学生活动数据
        '社会实践相关信息': ['实践起止时间','实践结束时间','实践团名称','实践主题','实践地点','实践团等级','指导老师','志愿者服务总时长','志愿服务项目名称','志愿服务日期','志愿服务地点','志愿服务时长','志愿服务组织名称','企业导师','基地名称','实践成绩','实践起始时间','开始时间','结束时间','志愿服务总时长'],
        # 竞赛数据
        '竞赛报名信息': ['队伍名称','学生身份','竞赛指导教师姓名','竞赛指导教师工作证号','参赛作品名称','指导老师','团队成员人数','竞赛形式'],
    },
    '教学域': {
        # 教学数据
        '考试信息': ['考号','外国语成绩','政治理论成绩','业务课成绩','加试成绩','加试科目','初试总成绩','复试成绩','初试权重','复试权重','总成绩','开课学期','课程编号','课程名称','成绩','成绩标识','学分','总学时','考核方式','考试性质','任课教师'],
        '教学成果信息': ['教学成果编号','教学成果名称','奖励级别码','授奖等级码','获奖年度','完成人总数'],
        # zy
        '教学经费信息': ['项目名称','金额'],
        # 进修学习相关信息
        '国内进修': ['进修去向','进修生生源','CSC登记号','所属项目','来华日期','入校日期','预计离校日期','毕业日期','在校状况码'],
        '出国（境）进修': ['批件号','批准天数','出境日期','入境日期','执行情况','经费来源','邀请单位','访问国家/地区','出访成果','出访任务','护照号', '通行证号','签发日期','失效日期'],
    },
    '科研域': {
        # 科研数据
        # zy
        '私密论文信息': ['论文题目','年份','发表时间','检索号','检索类型','刊物名称','卷','期','页码','是否ESI高被引论文','是否通讯作者','中科院JCR分区','影响因子','他引次数','是否第一作者'],
        '私密科研成果信息': ['项目名称','项目负责人','项目编号','立项日期','开始日期','计划结项日期','结项日期','项目性质','项目分类','位次','项目完成情况','项目来源','著作名称','书号','类别','出版时间','出版单位','本人位次','是否主编','本人撰写字数（千字）','总字数（千字）','所有作者'],
        '科研经费信息': ['经费来源','经费金额'],
        '科研合同信息': ['合同编号','合同名称','项目编号','甲方','乙方','两方','合同类型码','合同签订日期','合同截止日期','合同金额','是否已归档','签订地点合同起始日期','合同终止日期','课题分类','免证营业税','日期认定号','认定日期'],
        '科研机构相关信息': ['单位号','审批机构','审批日期','成立日期','管理体制码','定编人数','实有人数','业务管理人员数','用房面积','固定资产','从事研究的学科领域码','机构类型码','组成类型码','研究方向','办公地址','电子信箱','负责人职工号','负责人电话','负责人姓名','机构简介','机构名称','联系电话','主管部门编'],
        '技术转让': ['合同编号','合同名称','成交金额','本年实际收入','单位号','学科领域码','签定日期','负责人职工号','受让方名称','受让方类型码','社会经济效益码','所属项目编号'],
        # 学术交流
        '学术交流会议': ['会议中文名称','会议英文名称','会议地点','会议起始日期','会议终止日期','学科领域码','会议等级码','会议举办形式码','会议中文名称','会议英文名称','会议地点','会议起始日期','会议终止日期','学科领域码','会议等级码','会议举办形式码'],
        '派出人员': ['派出人员编号','交流类型码','学科领域码','国别码','交流开始日期','交流结束日期'],
        '接受人员': ['接受人员姓名','性别码','国别码','交流类型码','学科领域码','交流开始日期','交流结束日期'],
        # 项目数据
        # zy
        '私密项目信息': ['项目名称','项目负责人','项目编号','立项日期','开始日期','计划结项日期','结项日期','项目性质','项目分类','位次','项目完成情况','项目来源'],
    }
}

schemas_dict_education_D = {
    '人员域': {
        '学生身份信息': list({
            # 基本信息
            '姓名','姓名拼音','曾用名','民族','性别名称','港澳台侨','国家地区','手机号','电子信箱','微信号','QQ号','通讯地址','邮政编码',
            # 身份信息
            '身份证件类型','身份证件号',
            # 体质健康信息
            '身高等级','体重等级','肺活量等级','仰卧起坐等级','跑步等级','视力等级','篮球等级','足球等级',
            # 家庭成员信息
            '工作单位','家庭电话','家庭住址','家庭邮政编码','家庭关系','家庭成员姓名','联系方式','家长出生日期',
            # 资助信息
            '资助类型','资助年份','资助金额','资助学年','银行卡号','开户行',
            # 家庭经济信息
            '劳动力人口数','年人均收入','是否困难户',
            # zy
            # 网络密码及关联信息
            '系统登录账号','密码',
            }),
        '学生校内信息': list({
            # 学生入学信息
            '入学专业','学科门类','学习方式','培养方式','学生类别','是否在职生','在职工号','入学年月','入学年级','现在专业','研究方向','学制','学位类型','学籍状态','是否休学','是否请假','预计毕业年份','预计毕业月份','实际毕业时间',
            # 学籍信息
            '学籍相关信息','学籍异动信息','毕业中学','生源地',
            # 学历学位信息
            '学历码','所学专业','入学年月','学习形式码学制','结束学业年月','毕肄业学校或单位','毕业证号','毕业证书印制号','获得学位码',
            # 档案信息 
            '异动日期','异动文号','异动类别','异动原因描述','复学日期','原系所','原班级','原专业','离校状态','档案接受地址','档案去向地','档案邮寄地址','邮寄方式','运单号','原导师','档案邮寄时间',
            # 处分信息
            '违纪类型','违纪日期','处分类型','处分日期','处分原因','处分文号','处分解除日期','处分撤销日期','处分撤销原因','申诉日期','申委会审议日期','申委会审议结论','违纪类别','审委会审议结论',
            }),
        '教师身份信息': list({
            # 基本信息
            '姓名','性别','证件号码','出生日期','民族','最高学历','最高学位层次','获得最高学历的院校或机构',
            # 资格信息
            '资格种类','任教学科','证书号码','认证机构名称','专业技术职务名称','专业技术职务级别','取得资格时间','评审类型','评审学科','是否现职称',
            # 身份信息
            '学工号','学院','单位','专业','班级','导师工号','导师类别','导师类别代码','导师联系方式',
            # 工资待遇信息
            '年工资收入','基本工资','绩效工资','乡村教师生活补助','其他津贴补贴','五险一金',
            # zy
            # 网络密码及关联信息 
            '系统登录账号','密码',
            }),
        '干部身份信息': list({
            # 基本信息
            '姓名','性别','证件号码','出生日期','民族','最高学历','最高学位层次','获得最高学历的院校或机构',
            }),
    },
    '财务域': {
        '学生财务相关': list({
            # 一卡通消费信息
            '操作时间','科目描述','钱包交易金额','钱包余额','操作员','工作站','终端名称',
            # 创新基金信息
            '基金级别','资助年限','资助金额','研究成果'
            }),
        '账务管理': list({
            # zy
            # 账务数据
            '类型编号','类型名称','凭证标记','限制科目','限制性质码',
            # 凭证数据
            '凭证编号','分录标记码','凭证标记','凭证摘要','科目编号','部门代码','项目编号','记帐方向码','金额','生成日期','冲暂','支票号','备注','制单','复核','出纳',
            # 经费来源
            '会计年度','入帐月份','凭证标志','凭证日期',
            # 往来账数据
            '金额','原因',
            # 暂存款
            '金额','原因',
            # 暂付款
            '金额','原因',
            # 税务信息
            '金额','税务类型','抬头',
            # 票据信息
            '票本序号','启始号','截止号','票类码','规格','用途','领用人','联系方法','领用部门号','入帐科目及项目','金额限制','领用日期','归还日期','已用截止号','实开金额','有效张数','作废张数','交回人号','验收人号','备注',
            }),
    },
    '资产域': {
        '实验室设施数据': list({
            # 实验室资源数据
            '实验室号','实验室名称','建立日期','建筑面积','使用面积','实验室位置','负责人教工号','管理级别码','实验室类别码','实验室电话','器材管理员教工号','实验室介绍',
            # 实验室仪器
            '仪器号','仪器名称','英文名称','学校单位层次码','产权及使用状况码','分类号码型号','出厂日期','生产国别码','经费科目码','购置日期','厂家','出厂号','单据号',
            # zy
            # 实验室运行
            '实验室号','年度','投入额支出额','年耗电量','年耗水量','维护费',
            }),
    },
    '校园生活': {
        '学生活动数据': list({
            # 社会实践相关信息
            '实践起止时间','实践结束时间','实践团名称','实践主题','实践地点','实践团等级','指导老师','志愿者服务总时长','志愿服务项目名称','志愿服务日期','志愿服务地点','志愿服务时长','志愿服务组织名称','企业导师','基地名称','实践成绩','实践起始时间','开始时间','结束时间','志愿服务总时长',
            }),
        '竞赛数据': list({
            # 竞赛报名信息
            '队伍名称','学生身份','竞赛指导教师姓名','竞赛指导教师工作证号','参赛作品名称','指导老师','团队成员人数','竞赛形式',
            }),
    },
    '教学域': {
        '考试信息': list({
            # 考试信息
            '考号','外国语成绩','政治理论成绩','业务课成绩','加试成绩','加试科目','初试总成绩','复试成绩','初试权重','复试权重','总成绩','开课学期','课程编号','课程名称','成绩','成绩标识','学分','总学时','考核方式','考试性质','任课教师',
            # 教学成果信息
            '教学成果编号','教学成果名称','奖励级别码','授奖等级码','获奖年度','完成人总数',
            # zy
            # 教学经费信息
            '项目名称','金额',
            }),
        '进修学习相关信息': list({
            # 国内进修
            '进修去向','进修生生源','CSC登记号','所属项目','来华日期','入校日期','预计离校日期','毕业日期','在校状况码',
            # 出国（境）进修
            '批件号','批准天数','出境日期','入境日期','执行情况','经费来源','邀请单位','访问国家/地区','出访成果','出访任务','护照号', '通行证号','签发日期','失效日期',
            }),
    },
    '科研域': {
        '科研数据': list({
            # zy
            # 私密论文信息
            '论文题目','年份','发表时间','检索号','检索类型','刊物名称','卷','期','页码','是否ESI高被引论文','是否通讯作者','中科院JCR分区','影响因子','他引次数','是否第一作者',
            # 私密科研成果信息
            '项目名称','项目负责人','项目编号','立项日期','开始日期','计划结项日期','结项日期','项目性质','项目分类','位次','项目完成情况','项目来源','著作名称','书号','类别','出版时间','出版单位','本人位次','是否主编','本人撰写字数（千字）','总字数（千字）','所有作者',
            # 科研经费信息
            '经费来源','经费金额',
            # 科研合同信息
            '合同编号','合同名称','项目编号','甲方','乙方','两方','合同类型码','合同签订日期','合同截止日期','合同金额','是否已归档','签订地点合同起始日期','合同终止日期','课题分类','免证营业税','日期认定号','认定日期',
            # 科研机构相关信息
            '单位号','审批机构','审批日期','成立日期','管理体制码','定编人数','实有人数','业务管理人员数','用房面积','固定资产','从事研究的学科领域码','机构类型码','组成类型码','研究方向','办公地址','电子信箱','负责人职工号','负责人电话','负责人姓名','机构简介','机构名称','联系电话','主管部门编',
            # 技术转让
            '合同编号','合同名称','成交金额','本年实际收入','单位号','学科领域码','签定日期','负责人职工号','受让方名称','受让方类型码','社会经济效益码','所属项目编号',
        }),
        '学术交流': list({
            # 学术交流会议
            '会议中文名称','会议英文名称','会议地点','会议起始日期','会议终止日期','学科领域码','会议等级码','会议举办形式码','会议中文名称','会议英文名称','会议地点','会议起始日期','会议终止日期','学科领域码','会议等级码','会议举办形式码',
            # 派出人员
            '派出人员编号','交流类型码','学科领域码','国别码','交流开始日期','交流结束日期'
            # 接受人员
            '接受人员姓名','性别码','国别码','交流类型码','学科领域码','交流开始日期','交流结束日期',
            }),
        '项目数据': list({
            # zy
            # 私密项目信息
            '项目名称','项目负责人','项目编号','立项日期','开始日期','计划结项日期','结项日期','项目性质','项目分类','位次','项目完成情况','项目来源',
            }),
    },
}

def judge_text_pdf(file_path):
    """
        判断是否为文字类pdf。（即是否可以直接使用pdfplumber读取文字内容）
    :param file_path: 文件路径
    :return: True 或 False
    """
    doc = pdfplumber.open(file_path)
    for page in doc.pages:
        if page.extract_words():
            return True
    return False


def parse_pdf(file_path, max_page=-1):
    """
        文字类pdf内容提取。
    :param file_path: 文件路径
    :param max_page: 最大抽取页数，-1时抽取全部
    :return: 文本内容，str
    """
    pdf = pdfplumber.open(file_path)
    texts = list()
    for i, page in enumerate(pdf.pages):
        if 0 <= max_page <= i:
            break
        texts.append(page.extract_text())
    texts = '\n'.join(texts)
    pdf.close()
    return texts


def parse_pdf_with_ocr(file_path, ocr, max_page=-1):
    """
        扫描件类pdf内容提取。
    :param file_path: 文件路径
    :param ocr: OCR模型
    :param max_page: 最大抽取页数，-1时抽取全部
    :return: 文本内容，str
    """
    pdfDoc = fitz.open(file_path)
    if pdfDoc.page_count == 0:
        return ''

    mat = fitz.Matrix(1.33, 1.33)   # 图片缩放倍率，当前为1.33倍（此时图片的大小与wps中pdf转图片大小一致）
    texts = []
    for i, page in enumerate(pdfDoc):
        if 0 <= max_page <= i:
            break

        # 转图片
        pix = page.get_pixmap(matrix=mat, alpha=False)
        bytesdata = pix.tobytes('png')
        imgarray = np.frombuffer(bytesdata, dtype=np.uint8)
        img = cv2.imdecode(imgarray, cv2.IMREAD_ANYCOLOR)

        # 从图片中提取文字
        result = ocr.ocr(img, cls=False)
        result = result[0]
        for res in result:
            if res[1][0] != '':
                texts.append(res[1][0])
    texts = '\n'.join(texts)
    return texts


def parse_docx(file_path, ocr=None, parse_picture=False):
    """
        docx文件内容提取，提取文字段落、表格内容、图片ocr文本（可选）
    :param file_path: 文件路径
    :param ocr: ocr模型
    :param parse_picture: 是否对文件中的图片进行ocr
    :return: 提取出的文本内容，str
    """
    doc = docx.Document(file_path)
    texts = list()
    # 遍历文档中的段落
    for para in doc.paragraphs:
        texts.append(para.text)

    # 遍历文档中的表格
    for table in doc.tables:
        # 遍历表格中的行
        for row in table.rows:
            # 遍历行中的单元格
            for cell in row.cells:
                texts.append(cell.text)

    if parse_picture and not ocr is None:
        # 遍历文档中的图片
        dict_rel = doc.part._rels  # rels其实是个目录
        for rel in dict_rel:
            rel = dict_rel[rel]
            if "image" in rel.target_ref:
                bytesdata = rel.target_part.blob
                imgarray = np.frombuffer(bytesdata, dtype=np.uint8)
                img = cv2.imdecode(imgarray, cv2.IMREAD_ANYCOLOR)
                # 从图片中提取文字
                result = ocr.ocr(img, cls=False)
                result = result[0]
                for res in result:
                    if res[1][0] != '':
                        texts.append(res[1][0])

    texts = '\n'.join(texts)
    return texts


def get_text(file_path, ocr, max_page=-1):
    """
        获取pdf/docx文字内容。
    :param file_path: 文件路径
    :param ocr: OCR模型
    :param max_page: 最大抽取页数。-1时全部抽取
    :return: 文本内容，str
    """
    content = ''
    try:
        if judge_pdf(file_path):
            if judge_text_pdf(file_path):
                content = parse_pdf(file_path, max_page=max_page)
            else:
                content = parse_pdf_with_ocr(file_path, ocr, max_page=max_page)
        else:
            content = parse_docx(file_path, ocr, parse_picture=False)
    except Exception as e:
        print(pathlib.Path(file_path).stem + '[Failed Get Content] : ' + str(e))

    return content.replace('\n', '')


def pdf2img(file_path, dest_path):
    """
        pdf转图片。
    :param file_path: pdf文件路径
    :param dest_path: 图片存储路径
    :return: 所有图片存储路径列表，list
    """
    zoom_x = 2.0
    zoom_y = 2.0
    mat = fitz.Matrix(zoom_x, zoom_y)
    doc = fitz.open(file_path)
    image_paths = list()
    for page in doc:
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img_path = '%s_%s.png' % (os.path.basename(dest_path), page.number + 1)
        pix.save(os.path.join(dest_path, img_path))
        image_paths.append(os.path.join(dest_path, img_path))
    return image_paths


def get_text_info_from_single_image(img_path, ocr):
    """
        从单张图片中提取文字。
    :param img_path: 单张图片路径
    :param ocr: OCR模型
    :return: 文本内容，str；提取文本内容的详细信息，json
    """
    img = cv2.imread(img_path)

    result = ocr.ocr(img, cls=False)
    json_data = []
    for text_block in result[0]:
        json_data.append({
            'position': text_block[0],
            'text': text_block[1][0],
            'confidence': text_block[1][1]
        })

    json_data = json.dumps(json_data, ensure_ascii=False)
    texts = '\n'.join([it['text'] for it in json.loads(json_data)])
    return texts, json_data


def get_embedding_table(embedding_path):
    """
        读入词嵌入表（腾讯词向量精简版，10万词，200维）
    :param embedding_path: 词嵌入表路径
    :return: 词嵌入表，dict
    """
    embeddings = {}
    mean_vec = [0.0] * 200  # 默认值，当遇到OOV词时，初始化为该向量
    with open(embedding_path, 'r', encoding='utf-8') as f:
        f.readline()
        for line in f.readlines():
            line = line.strip().split()
            word = line[0]
            embedding = list(map(lambda x: float(x), line[1:]))
            embeddings[word] = embedding

    embeddings['mean_vector'] = mean_vec

    return embeddings


def get_embedding(word, embeddings):
    """
        获取特定词向量。若匹配到了指定词，返回对应的词向量，否则，返回默认词向量（全0向量）。
    :param word: 待获取词向量的词，str
    :param embeddings: 词嵌入表，dict
    :return: 词嵌入的向量，list；是否匹配到了该词，bool；
    """
    if word in embeddings:
        return embeddings[word], True  # matched = True
    else:
        return embeddings['mean_vector'], False  # matched = False


def get_domain_keywords(keywords_path):
    """
        读入领域关键词字典。（格式：{领域：{关键词: [词向量]}}）
    :param keywords_path: 领域关键词表路径
    :return: 领域关键词表，dict
    """
    domain_keywords = {}
    with open(keywords_path, 'r', encoding='utf-8') as f:
        domain_keywords = eval(f.read())
    return domain_keywords


def get_intersection(intersection_path):
    """
        读入领域关键词交集（格式：{关键词}）
    :param intersection_path: 交集词表路径
    :return: 交集词表，set
    """
    intersection = set()
    with open(intersection_path, 'r') as f:
        intersection = eval(f.read())
    return intersection


def have_chinese(text):
    """
        判断文件名/标题文字中是否包含中文字符。
    :param text: 文件名/标题，str
    :return: 是否包含中文，bool
    """
    for c in text:
        if '\u4e00' <= c <= '\u9fff':
            return True
    return False


def judge_pdf(file_path):
    """
        判断一个文件是不是pdf文件
    :param file_path: 文件路径，str
    :return: 是否是pdf文件，bool
    """
    return file_path.lower().endswith(".pdf")


def get_title(file_path, ocr, table_engine):
    """
        根据文件类型获取文件的标题
    :param file_path: 文件路径
    :param ocr: OCR模型
    :param table_engine: 版面识别模型
    :return: 标题，str
    """
    if judge_pdf(file_path):
        return get_pdf_title(file_path, ocr, table_engine)
    else:
        return get_docx_title(file_path)


def get_docx_title(file_path):
    """
        把docx文件的第一个文字段当作标题
    :param file_path: 文件路径
    :return: 文件的标题，str
    """
    doc = docx.Document(file_path)
    for para in doc.paragraphs:
        if len(para.text.strip()) > 0:
            return para.text.strip()[:15]
    return ""


def get_pdf_title(file_path, ocr, table_engine):
    """
        从文件首页提取标题。
    :param file_path: pdf文件路径
    :param ocr: OCR模型
    :param table_engine: 版面识别模型
    :return: 标题，str
    """
    ## 打开文件
    pdfDoc = fitz.open(file_path)
    if pdfDoc.page_count == 0:
        return ''

    ## 首页转图片
    page = pdfDoc[0]
    mat = fitz.Matrix(1.33, 1.33)
    pix = page.get_pixmap(matrix=mat, alpha=False)

    bytesdata = pix.tobytes('png')
    imgarray = np.frombuffer(bytesdata, dtype=np.uint8)
    img = cv2.imdecode(imgarray, cv2.IMREAD_ANYCOLOR)

    ## 版面识别
    result = table_engine(img)

    ## 标题提取
    labels = ['title', 'table_caption']
    titles = []
    for res in result:
        if res["type"] in labels:
            img = res['img']
            title = ''
            pad = 15
            img = cv2.copyMakeBorder(img, pad, pad, pad, pad, cv2.BORDER_CONSTANT, value=[255, 255, 255])

            ocrres = ocr.ocr(img, cls=True)
            ocrres = ocrres[0]
            for r in ocrres:
                if r[1][0] != '':
                    title += r[1][0]
            titles.append(title)

    return ' '.join(titles).strip()


def get_filename(file_path, ocr, table_engine):
    """
        获取文件名/标题：当文件名包含中文时，返回文件名；否则，从文件内抽取标题，返回标题。
    :param file_path: pdf/docx文件路径
    :param ocr: OCR模型
    :param table_engine: 版面识别模型
    :return:原文件名，新文件名/标题，str
    """
    ori_filename = pathlib.Path(file_path).stem
    filename = ori_filename
    invalidstr = r"[\/\\\:\*\?\"\<\>\|\0]"
    try:
        if not have_chinese(filename):
            title = get_title(file_path, ocr, table_engine)
            if len(title) > 0:
                filename = re.sub(invalidstr, '', title)
    except Exception as e:
        print('Extract Title Error: [{}]'.format(filename) + str(e))

    return ori_filename, filename.strip().replace('\n', '')


def shorten_filename(filename, limit=100):
    """
        截断文件名，防止过长。
    :param filename: 原文件名，str
    :param limit: 限制的最大长度，默认为100字符
    :return: 新文件名，str
    """
    if len(filename) <= limit:
        return filename
    else:
        return filename[:limit]


def get_sensitive_level(text, field, is_title=True):
    """
        基于敏感词匹配的敏感分级。
    :param text: 文件名/标题/文本内容
    :param field: 领域
    :param is_title: 是否为标题，默认为True
    :return: 敏感/非敏感，当领域未知时，返回NaN，str
    """
    if field == 'industry':
        sensitive_words_title = {'标书', '合同', '发票'}
        sensitive_words_content = {'秘密', '机密', '绝密', '不允许公开'}

        if is_title:
            for i in sensitive_words_title:
                if i in text:
                    return '敏感'
        else:
            for i in sensitive_words_content:
                if i in text:
                    return '敏感'
        return '非敏感'
    else:
        return 'NaN'


def get_classification(text, field):
    """
        基于分级分类关键词的文件分类。
    :param text: 文件名/标题
    :param field: 领域
    :return: 分类类别/NaN，str
    """
    if field == 'industry':
        class_dict = {'A1-3:设计图纸数据': {'设计图', '平面图', '施工图', '原理图', '详图', '宗地图', '总平图',
                                      '剖面图', '断面图', '构造图', '立面图', '设计图', '设计报告', '建筑设计', '平面设计'},
                      'C1-1:物流数据': {'托运人', '收获人', '装货人', '承运人', '目的地', '起运地', '起运日期',
                                    '驾驶员', '押运员', '车牌号码', '道路运行证', '调度人'},
                      'D1:设备资产域数据': {'总价', '价格', '价税合计'},
                      'D2:产品供应链数据': {'实施机构', '建设单位', '设计单位', '承运人', '委托方', '生产方', '使用单位',
                                     '测试单位', '勘察单位', '企业名称', '责任单位'},
                      'P1-1:用户资本资料': {'姓名', '性别', '民族', '年龄', '住址', '身份证号', '毕业日期', '毕业院校',
                                      '学历类别', '学制', '学习形式', '工作单位/聘用企业', '注册专业', '任职资格'},
                      'P1-2:用户身份证明': {'身份证', '毕业证', '学位证', '从业资格证', '建造师注册证书'}}
        for k, v in class_dict.items():
            for word in v:
                if word in text:
                    return k
        return 'NaN'
    else:
        return 'NaN'


def get_total_schema_length(schemas):
    """
        统计某领域下schema词典中总的关键词个数。
    :param schemas: 某领域下的schema词典，dict
    :return: 总关键词个数，int
    """
    total_length = 0
    for single_schema in schemas.values():
        total_length += len(single_schema)
    return total_length


def get_file_MG_level(extraction_result_list, total_schema_length):
    """
        根据抽取的schema内容数，确定敏感等级。（最初基于抽取数/总schema关键词个数的比值，但有效性有待商榷，目前暂定为只根据抽取数）
    :param extraction_result_list: UIE信息抽取结果，list
    :param total_schema_length: 对应领域下总schema关键词个数（基于比例确定敏感等级时使用），int
    :return: 不敏感/低敏感/较敏感/敏感，str
    """
    extraction_result_list_length = len(extraction_result_list)
    loop = 0
    extraction_num = 0
    while loop < extraction_result_list_length:
        if loop % 2 != 0:
            single_schema_extraction_results = extraction_result_list[loop].values()
            for ele in single_schema_extraction_results:
                extraction_num += len(ele)
        loop += 1
    ratio = (float(extraction_num) / float(total_schema_length)) if total_schema_length != 0 else 0
    ret_MG_level = ''
    # if ratio <= 0.25:  # 比例法有待商榷，schema个数差距太大
    #     ret_MG_level = '不敏感'
    # elif ratio <= 0.5:
    #     ret_MG_level = '低敏感'
    # elif ratio <= 0.75:
    #     ret_MG_level = '较敏感'
    # else:
    #     ret_MG_level = '敏感'

    if extraction_num < 5:
        ret_MG_level = '不敏感'
    elif extraction_num < 10:
        ret_MG_level = '低敏感'
    elif extraction_num < 20:
        ret_MG_level = '较敏感'
    else:
        ret_MG_level = '敏感'

    return ret_MG_level


def get_pdf_table(file_path, save_dir):
    """
        识别并提取表格到excel中。（pdf文件中目前只支持文字类pdf的表格抽取）
    :param file_path: pdf文件路径
    :param save_dir: excel存储路径
    :return: 无
    """
    try:
        pdf = pdfplumber.open(file_path)
        wb = Workbook()  # 实例化一个工作簿对象
        ws = wb.active  # 获取第一个sheet
        # 获取每一页的表格中的文字，返回table、row、cell格式：[[[row1],[row2]]]
        for page in pdf.pages:
            for table in page.extract_tables():
                for row in table:
                    # 对每个单元格的字符进行简单清洗处理
                    row_list = [cell.replace(
                        '\n', ' ') if cell else '' for cell in row]
                    ws.append(row_list)  # 写入数据
        wb.save(os.path.join(save_dir, pathlib.Path(file_path).stem + '.xlsx'))
        pdf.close()
    except Exception as e:
        print('Table extraction failed!', e)


def read_file(file_path, ocr, table_engine, table_extract=False, table_dir=None):
    """
        读入pdf/docx文件，获取文字内容
    :param file_path: 文件绝对路径
    :param ocr: OCR模型
    :param table_engine: 版面识别模型
    :param table_extract: 是否提取表格到excel
    :param table_dir: 表格存储目录
    :return: 原始文件名，str；提取文件名，str；文字读取结果，str
    """

    if not os.path.exists(file_path):
        raise ValueError('file path is not exist!', file_path)
    if not os.path.isfile(file_path):
        raise ValueError('not a correct file path!', file_path)
    if not file_path.lower().endswith(".pdf") and not file_path.lower().endswith(".docx"):
        raise ValueError('not a supported file type! (pdf/docx)', file_path)

    # 获取无标题文件的标题
    ori_filename, filename = get_filename(file_path, ocr, table_engine)  # 不包含后缀，且删除非法字符的文件名

    # 获取文字内容
    content = get_text(file_path, ocr)

    # 单独提取表格
    if table_extract:
        extract_table(file_path, table_dir)

    return ori_filename, filename, content


def get_docx_table(file_path, save_dir):
    """
        识别并提取表格到excel中。
    :param file_path: 文件路径
    :param save_dir: excel存储路径
    :return: 无
    """
    try:
        doc = docx.Document(file_path)
        wb = Workbook()  # 建一个新的工作簿
        sht = wb.active  # 获取第一个sheet
        for table in doc.tables:
            # 获取表格的行
            tb_rows = table.rows
            # 读取每一行内容
            for i in range(len(tb_rows)):
                row_data = []
                row_cells = tb_rows[i].cells
                # 读取每一行单元格内容
                for cell in row_cells:
                    row_data.append(cell.text)  # 单元格内容
                sht.append(row_data)  # 逐行写到excel中
        wb.save(os.path.join(save_dir, pathlib.Path(file_path).stem + '.xlsx'))
    except Exception as e:
        print('Table extraction failed!', e)


def extract_table(file_path, table_dir):
    """
        单独提取文件中的表格
    :param file_path: 文件路径
    :param table_dir: 存储的表格路径
    :return: 无
    """
    os.makedirs(table_dir, exist_ok=True)
    if judge_pdf(file_path):
        if judge_text_pdf(file_path):
            get_pdf_table(file_path, table_dir)
    else:
        get_docx_table(file_path, table_dir)


def get_domain_type(text, embeddings, domain_keywords, intersection):
    """
        基于文件名/标题文本的领域分类。
    :param text: 文件名/标题
    :param embeddings: 词嵌入表
    :param domain_keywords: 领域关键词
    :param intersection: 领域关键词交集
    :return:
        domain_pred_res: 领域预测结果，str
        simi_sum / match_num: 预测相似度/置信度，float
        words: 文件名/标题关键词提取结果，set
        domain_matched: 各领域下匹配到的文本关键词，dict {领域: [文本关键词, 匹配的最佳领域词, 相似度]}
    """
    # 分词，过滤，删交集
    words = set()
    for word in jieba.lcut(text, cut_all=True):
        if len(word) > 1:
            words.add(word)
    words = words - intersection

    # 计算与各领域的匹配情况
    # 每个领域都get最大的(一个或多个)，然后保留最大的一个或多个作为匹配结果
    domain_pred = {}  # 领域：[匹配个数，总相似度]
    domain_matched = {}  # 领域：[匹配到的关键词]
    for domain in domain_keywords.keys():
        domain_pred[domain] = [0, 0]
        domain_matched[domain] = []

    for word in words:
        domain_max = {}
        for k in domain_keywords.keys():
            domain_max[k] = {'max_simi': -1, 'max_word': []}

        # 首先尝试直接完全匹配
        fully_matched = False
        for domain, keywords in domain_keywords.items():
            if word in keywords:
                domain_max[domain]['max_simi'] = 1
                domain_max[domain]['max_word'] = [word]
                fully_matched = True
                continue

        # 若完全匹配上，则其他领域也仅需判断是否可完全匹配即可，无需计算其他的相似度
        # 否则，计算与各领域关键词的相似度
        if not fully_matched:
            word_vec, matched = get_embedding(word, embeddings)
            if matched:  # 可向量化时，再继续计算相似度，否则跳过
                for domain, keywords in domain_keywords.items():
                    for k, v in keywords.items():
                        simi = \
                        F.cosine_similarity(paddle.to_tensor(word_vec), paddle.to_tensor(v), axis=0).tolist()[0]
                        if simi > domain_max[domain]['max_simi']:
                            domain_max[domain]['max_simi'] = simi
                            domain_max[domain]['max_word'] = [k]
                        elif simi == domain_max[domain]['max_simi']:
                            domain_max[domain]['max_word'].append(k)

        # 获取最大相似度
        max_simi = -1
        for v in domain_max.values():
            if v['max_simi'] > max_simi:
                max_simi = v['max_simi']

        # 获取最终匹配结果
        for k, v in domain_max.items():
            if v['max_simi'] > 0.5 and v['max_simi'] == max_simi:  # 阈值设为0.5
                domain_pred[k][0] += len(v['max_word'])
                domain_pred[k][1] += (len(v['max_word']) * max_simi)
                for w in v['max_word']:
                    domain_matched[k].append([word, w, max_simi])

    # 确定类别。匹配个数多者优先，相同时相似度大者优先，各领域均为0时，返回“others”
    domain_pred_res = 'others'
    match_num = -1
    simi_sum = -1
    for k, v in domain_pred.items():
        if v[0] == 0:
            continue
        if v[0] > match_num:
            match_num = v[0]
            simi_sum = v[1]
            domain_pred_res = k
        elif v[0] == match_num and v[1] > simi_sum:
            simi_sum = v[1]
            domain_pred_res = k

    return domain_pred_res, simi_sum / match_num, words, domain_matched


def extract_information(file_name, file_txt, log_base_dir, field, schemas, uie_dict, ori_filename=None):
    """
        提取文件关键信息，完成领域分类、敏感信息提取、敏感分级分类。
    :param file_name: 文件名
    :param file_txt: 文件文字内容
    :param log_base_dir: 结果文件存储路径（文件夹）
    :param field: 领域类型
    :param schemas: 特定领域下的schema词典
    :param uie_dict: 存储了预加载UIE模型的字典
    :param ori_filename: 原始文件名，用于log文件的文件名。若无，则使用file_name
    :return: 输出log文件，并返回基于schema抽取结果的敏感分级结果和是否包含机密词
    """
    try:
        # 基于标题/文件名的信息获取
        filename_mg_level = get_sensitive_level(file_name, field=field)
        filedata_fenji = get_classification(file_name, field=field)

        os.makedirs(log_base_dir, exist_ok=True)

        if ori_filename is not None:
            out_log_path = os.path.join(log_base_dir, shorten_filename(ori_filename) + '.log')
        else:
            out_log_path = os.path.join(log_base_dir, shorten_filename(file_name) + '.log')
        out_log_path_file = open(out_log_path, 'w')
        pprint('Processing : ' + file_name,
               stream=out_log_path_file)
        pprint('Domain_type : ' + field, stream=out_log_path_file)
        pprint('filename_MG_level : ' + filename_mg_level,
               stream=out_log_path_file)
        pprint('filedata_fenji : ' + filedata_fenji,
               stream=out_log_path_file)

        # 基于文件内容的信息抽取
        ## 机密词匹配
        have_sensitive_words = any((word in file_txt) for word in sensitive_words)
        if have_sensitive_words:
            pprint('是否包含机密词 : 是', stream=out_log_path_file)
        else:
            pprint('是否包含机密词 : 否', stream=out_log_path_file)

        ## schema信息抽取
        res = []
        if file_txt != '':
            for schema_name, schema_list in schemas.items():
                information_extraction = uie_dict[field][schema_name]
                ie_result = information_extraction(file_txt)  # 信息抽取
                res.append({'schema': schema_name})
                res.append(ie_result[0])
        file_MG_level = get_file_MG_level(res, get_total_schema_length(schemas))
        pprint('file_MG_level : ' + file_MG_level, stream=out_log_path_file)
        pprint(res, stream=out_log_path_file)
        out_log_path_file.close()

        if have_sensitive_words:
            return file_MG_level, '是'
        else:
            return file_MG_level, '否'
    except Exception as e:
        print(ori_filename + ' [Extraction Error] : ' + str(e))
        return None, None

def generate_ngram_lsh_fingerprint(text, n=1, num_perm=128, threshold=0.5):
    """
    使用N-gram + LSH算法生成文本的指纹（哈希值）。
    :param text: 要生成指纹的文本。
    :param n: N-gram的大小。
    :param num_perm: MinHash的排列数。
    :param threshold: LSH的阈值。
    :return: 文本的指纹（哈希值）。
    """
    # 创建MinHash对象
    minhash = MinHash(num_perm=num_perm)
    # 创建哈希函数（可以选择MD5、SHA-1等哈希算法）
    hash_function = hashlib.md5
    # 生成N-gram序列
    ngram_sequence = ngrams(text, n)
    # 添加N-gram序列到MinHash中
    for ngram in ngram_sequence:
        for item in ngram:
            minhash.update(hash_function(item.encode()).digest())
    # 返回指纹的哈希摘要
    fingerprint_digest = minhash.digest()
    return fingerprint_digest


def hash_fingerprint(fingerprint):
    """
    使用哈希函数对指纹进行进一步的处理。
    :param fingerprint: 要处理的指纹。
    :return: 处理后的指纹（哈希值）。
    """
    # 创建哈希函数（可以选择MD5、SHA-1、SHA-256等哈希算法）
    hash_function = hashlib.md5
    # 计算指纹的哈希值
    hashed_fingerprint = hash_function(fingerprint).hexdigest()
    return hashed_fingerprint


def hamming_distance(fingerprint1, fingerprint2):
    """
    计算两个指纹之间的汉明距离。
    :param fingerprint1: 第一个指纹。
    :param fingerprint2: 第二个指纹。
    :return: 汉明距离。
    """
    distance = sum(c1 != c2 for c1, c2 in zip(fingerprint1, fingerprint2))
    return distance


def calculate_similarity(fingerprint1, fingerprint2):
    """
    计算两个指纹之间的相似度。
    :param fingerprint1: 第一个指纹。
    :param fingerprint2: 第二个指纹。
    :return: 相似度。
    """
    hamming_dist = hamming_distance(fingerprint1, fingerprint2)
    similarity = 1 - (hamming_dist / len(fingerprint1))
    return similarity
